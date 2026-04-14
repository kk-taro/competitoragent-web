"""
竞品分析 Agent — Web 版
底层：LangGraph ReAct Agent + OpenAI兼容接口（中转站）
功能：静态Demo报告（无需Key）+ 自定义实时分析 + Word下载
"""

import streamlit as st
import re
import textwrap
import requests
import io
from typing import Annotated, Literal
from langchain_core.messages import AIMessage, BaseMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langchain_core.tools import tool
from langgraph.graph import END, StateGraph
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode
from typing_extensions import TypedDict

# ── 页面配置 ──────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="竞品分析 Agent",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── 样式 ──────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
    --ink: #1a1a2e;
    --accent: #e94560;
    --accent2: #0f3460;
    --muted: #8892a4;
    --surface: #f8f7f4;
    --border: #e2ddd5;
    --gold: #c9a84c;
    --green: #7ec8a0;
}

html, body, [class*="css"] {
    font-family: 'Noto Serif SC', serif;
    background: var(--surface);
    color: var(--ink);
}

/* Hero */
.hero {
    background: var(--ink);
    color: white;
    padding: 2.5rem 2rem 2rem;
    margin: -1rem -1rem 2rem;
    border-bottom: 3px solid var(--accent);
    position: relative;
    overflow: hidden;
}
.hero::before {
    content: '';
    position: absolute;
    top: -50%; right: -10%;
    width: 400px; height: 400px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(233,69,96,0.15) 0%, transparent 70%);
}
.hero h1 { font-size: 2rem; font-weight: 700; margin: 0 0 0.3rem; letter-spacing: -0.5px; }
.hero .subtitle { color: var(--muted); font-size: 0.9rem; font-family: 'JetBrains Mono', monospace; }
.hero .api-hint { font-size: 0.78rem; color: rgba(255,255,255,0.55); font-family: 'JetBrains Mono', monospace; margin-bottom: 0.8rem; }
.hero .badge {
    display: inline-block; background: var(--accent); color: white;
    font-size: 0.7rem; padding: 2px 8px; border-radius: 2px;
    font-family: 'JetBrains Mono', monospace; margin-left: 8px; vertical-align: middle;
}

/* Demo 卡片 */
.demo-card {
    background: white; border: 1px solid var(--border);
    border-left: 4px solid var(--accent); border-radius: 6px;
    padding: 1.2rem 1.4rem; margin-bottom: 0.4rem;
    transition: all 0.2s; position: relative;
}
.demo-card:hover { border-left-color: var(--gold); box-shadow: 0 4px 16px rgba(0,0,0,0.1); transform: translateY(-1px); }
.demo-card.active { background: #fffbf0; border-left-color: var(--gold); }
.demo-card .track { font-size: 0.68rem; font-family: 'JetBrains Mono', monospace; color: var(--accent); text-transform: uppercase; letter-spacing: 1.5px; margin-bottom: 0.4rem; }
.demo-card .title { font-weight: 700; font-size: 1rem; color: var(--ink); margin-bottom: 0.3rem; }
.demo-card .status { font-size: 0.72rem; font-family: 'JetBrains Mono', monospace; color: var(--muted); margin-top: 0.3rem; }
.demo-card.active .status { color: #16a34a; font-weight: 600; }
/* 透明覆盖按钮 */
.card-wrap { position: relative; }
.card-wrap .stButton { position: absolute; top: 0; left: 0; width: 100%; height: 100%; opacity: 0; }
.card-wrap .stButton > button { width: 100%; height: 100%; cursor: pointer; }
}


/* 自定义分析区 */
.custom-box {
    background: white; border: 1px solid var(--border);
    border-radius: 6px; padding: 1.5rem 1.8rem;
    margin-top: 1.5rem;
}
.custom-box h4 { color: var(--ink); font-size: 1rem; margin: 0 0 1rem; }

/* Agent 日志 */
.log-box {
    background: #0d1117; color: #7ec8a0;
    font-family: 'JetBrains Mono', monospace; font-size: 0.78rem;
    padding: 1rem 1.2rem; border-radius: 4px;
    border-left: 3px solid var(--green);
    max-height: 280px; overflow-y: auto;
    line-height: 1.8; margin-bottom: 1rem;
    white-space: pre-wrap;
}

/* 报告正文 */
.report-body { line-height: 2; }
.report-body h1 { font-size: 1.6rem; color: var(--ink); border-bottom: 2px solid var(--accent); padding-bottom: 0.5rem; margin-bottom: 1.5rem; }
.report-body h2 { font-size: 1.15rem; color: var(--accent2); margin-top: 2rem; border-left: 3px solid var(--gold); padding-left: 0.8rem; }
.report-body h3 { font-size: 1rem; color: var(--ink); margin-top: 1.2rem; }
.report-body table { border-collapse: collapse; width: 100%; font-size: 0.875rem; margin: 1rem 0; }
.report-body th { background: var(--ink); color: white; padding: 8px 12px; text-align: left; }
.report-body td { border: 1px solid var(--border); padding: 7px 12px; }
.report-body tr:nth-child(even) td { background: #fafafa; }
.report-body a { color: var(--accent2); text-decoration: underline; }

/* 侧边栏 */
section[data-testid="stSidebar"] { background: var(--ink) !important; border-right: 1px solid rgba(255,255,255,0.08); }
section[data-testid="stSidebar"] * { color: white !important; }
section[data-testid="stSidebar"] .stTextInput input {
    background: rgba(255,255,255,0.92) !important;
    border: 1px solid rgba(255,255,255,0.4) !important;
    color: #1a1a2e !important; font-family: 'JetBrains Mono', monospace !important;
}
section[data-testid="stSidebar"] .stTextInput input::placeholder {
    color: #8892a4 !important;
}

/* 主区域输入框 */
.main .stTextInput input,
[data-testid="stTextInput"] input {
    background: white !important;
    border: 1px solid var(--border) !important;
    color: #1a1a2e !important;
}
.main .stTextInput input::placeholder,
[data-testid="stTextInput"] input::placeholder { color: #aab0bb !important; }
section[data-testid="stSidebar"] .stTextInput input,
section[data-testid="stSidebar"] [data-testid="stTextInput"] input {
    background: rgba(255,255,255,0.92) !important;
    border: 1px solid rgba(255,255,255,0.4) !important;
    color: #1a1a2e !important;
}
section[data-testid="stSidebar"] .stTextInput input::placeholder,
section[data-testid="stSidebar"] [data-testid="stTextInput"] input::placeholder {
    color: #8892a4 !important;
}


/* 按钮 */
.stButton > button {
    background: var(--accent) !important; color: white !important;
    border: none !important; border-radius: 4px !important;
    font-family: 'Noto Serif SC', serif !important; font-weight: 600 !important;
    padding: 0.5rem 1.5rem !important; transition: all 0.2s !important;
}
.stButton > button:hover { background: #c73652 !important; transform: translateY(-1px) !important; }

hr { border-color: var(--border); }

/* 标签页 */
.stTabs [data-baseweb="tab-list"] { gap: 8px; }
.stTabs [data-baseweb="tab"] {
    background: white; border: 1px solid var(--border);
    border-radius: 4px; padding: 6px 16px;
    font-family: 'Noto Serif SC', serif; font-size: 0.9rem;
}
.stTabs [aria-selected="true"] { background: var(--ink) !important; color: white !important; border-color: var(--ink) !important; }
</style>
""", unsafe_allow_html=True)

# ── 三份静态 Demo 报告 ────────────────────────────────────────────────────────

DEMO_REPORT_AI_CHAT = """# 竞品调研报告：国内AI对话助手

> **报告说明**：本报告基于公开信息整理，数据截至2025年Q1。来源包括官网、36kr、极客公园、AppSo等媒体。

## 一、报告概述（Executive Summary）

国内AI对话助手赛道在2024年进入白热化竞争阶段。豆包、Kimi、DeepSeek三家产品各具差异化定位：豆包依托字节生态打流量，Kimi以长文本技术能力见长，DeepSeek则以开源+极致性价比形成独特壁垒。整体市场呈现「技术趋同、生态分化」的格局，2025年竞争重心从模型能力转向场景深度和用户留存。

**核心结论：**
- 豆包DAU已超3000万（字节跳动财报，2024Q4），用户规模领先
- Kimi以200万+付费用户领跑变现能力（月之暗面融资文件，2024）
- DeepSeek开源策略重构行业格局，API调用成本仅为GPT-4的1/20

## 二、市场与赛道分析（Market Context）

**市场规模与增速**

- 2024年中国AI对话助手市场规模约180亿元，同比增长240%（IDC，2024年报）
- 预计2027年突破800亿，CAGR约65%（艾瑞咨询预测）
- 月活用户超1亿，是2023年同期的5倍（QuestMobile，2025Q1）

**竞争格局**

| 梯队 | 玩家 | 核心优势 |
|------|------|---------|
| 第一梯队 | 豆包、文心一言 | 流量生态 + 用户规模 |
| 第二梯队 | Kimi、通义 | 技术差异化 |
| 第三梯队 | DeepSeek、智谱 | 开源/B端 |

**趋势判断**
- 多模态能力（图像、语音、视频）成为标配，纯文本对话竞争力下降
- B端落地加速，企业知识库、客服、代码辅助是主要场景
- 开源模型冲击闭源定价体系，API价格战已蔓延至C端会员定价

## 三、竞品选择与分层（Competitive Landscape）

```
豆包    ──── 流量型 ──── 字节生态驱动，重分发轻技术
Kimi    ──── 技术型 ──── 长上下文差异化，重付费转化
DeepSeek ─── 开源型 ──── 成本优势+开发者生态，重API商业化
```

三款产品面向不同核心用户群：豆包主打泛用户（学生+白领），Kimi深耕知识工作者，DeepSeek以开发者和企业为主。

## 四、核心能力拆解（Product Capability Analysis）

### 豆包

- **产品定位**：字节系AI超级App，入口级产品，主打"日常使用的AI助手"
- **核心功能**：对话、写作、搜索、图片生成、语音对话、智能体（Coze接入）
- **技术特点**：基于字节自研Doubao模型，上下文128K，支持实时联网搜索
- **定价策略**：基础版免费；豆包Pro 19.9元/月，主推年付148元（来源：[豆包官网定价页](https://www.doubao.com)）
- **用户规模**：DAU超3000万，注册用户1.5亿+（字节2024Q4财报）
- **更新节奏**：每月1-2次大版本更新，迭代极快
- **分发渠道**：抖音、今日头条、西瓜视频内嵌入口，App独立下载，小程序
- **商业模式**：C端会员订阅 + B端API（火山引擎）+ 广告联动
- **近期动态**：2025年推出豆包1.5模型，推理能力接近GPT-4o；Coze国际版上线

### Kimi

- **产品定位**：以"长文本理解"为核心差异点的知识型AI助手
- **核心功能**：超长文档解析（支持200万Token）、联网搜索、代码生成、深度研究模式
- **技术特点**：Moonshot-v1系列模型，业界最长上下文窗口之一，RAG能力强
- **定价策略**：基础版免费（128K上下文）；Kimi+ 29.9元/月，年付199元（来源：[Kimi官网](https://kimi.moonshot.cn)）
- **用户规模**：注册用户5000万+，付费用户200万+（36kr报道，2024Q4）
- **更新节奏**：核心模型每季度迭代，功能更新每2-4周一次
- **分发渠道**：独立App、网页、浏览器插件（Chrome/Edge）、API
- **商业模式**：C端订阅（主要收入） + API商业化（企业客户）
- **近期动态**：2025年推出"深度研究"功能（类Perplexity），月之暗面估值达30亿美元

### DeepSeek

- **产品定位**：开源优先的AI基础模型公司，C端产品为模型能力的展示窗口
- **核心功能**：对话、代码生成、数学推理、长文本处理，主打"思维链"可视化
- **技术特点**：DeepSeek-V3/R1模型开源，MoE架构，训练成本仅600万美元（来源：[DeepSeek技术报告](https://arxiv.org/abs/2401.02954)）
- **定价策略**：C端App完全免费；API定价极低（输入0.1元/M tokens，为GPT-4的1/50）
- **用户规模**：App下载量超1亿（苹果应用商店，2025年1月曾登顶全球免费榜）
- **更新节奏**：模型大版本约每季度一次，开源社区持续迭代
- **分发渠道**：独立App、网页、完全开源（HuggingFace/GitHub）、API
- **商业模式**：主要依赖API商业化（B端）；C端免费作为品牌建设手段
- **近期动态**：DeepSeek-R1发布后席卷全球，在美国引发AI竞赛讨论；R2版本预计2025Q2发布

## 五、商业模式分析（Monetization）

| 维度 | 豆包 | Kimi | DeepSeek |
|------|------|------|---------|
| 主要收入来源 | 订阅 + API + 广告 | 订阅 + API | API（B端） |
| C端定价 | 19.9元/月 | 29.9元/月 | 免费 |
| API定价 | 中等（火山引擎） | 中等 | 极低 |
| 付费转化路径 | 免费→会员（功能限制触发） | 免费→会员（Token限制触发） | 基本不转化（C端） |
| 变现成熟度 | 中 | 高 | 低（C端）/高（B端） |

**关键洞察**：Kimi的付费转化率在三者中最高，核心原因是其用户（知识工作者）对200万Token上下文有刚性需求，愿意为此付费。

## 六、增长与分发策略（Growth Strategy）

**豆包**：依托字节流量矩阵，在抖音、今日头条等App内嵌AI入口，新用户获取成本接近于零；同时通过Coze平台吸引开发者构建智能体，形成生态飞轮。

**Kimi**：早期靠"PDF神器"口碑在知识工作者圈子裂变，后续通过浏览器插件占据用户日常工作场景；KOL种草（B站、小红书）是重要获客渠道。

**DeepSeek**：开源策略是最大增长引擎——开源模型被全球开发者下载部署，形成巨大自然流量；2025年1月凭借R1的惊艳表现登顶App Store，获得大量免费媒体曝光。

## 七、用户与场景分析（User & Use Case）

| 用户群 | 豆包 | Kimi | DeepSeek |
|--------|------|------|---------|
| 核心用户 | 学生、泛白领 | 研究者、知识工作者 | 开发者、工程师 |
| 年龄段 | 18-35岁为主 | 25-40岁为主 | 20-35岁为主 |
| 核心场景 | 日常问答、创意写作、娱乐 | 文档分析、研究报告、学习 | 代码辅助、数学推理、技术问答 |
| 使用频率 | 高频（日常） | 中高频（工作日） | 中频（按需） |

## 八、优劣势对比（SWOT / 对比矩阵）

### 豆包 SWOT

| | 内容 |
|--|------|
| **优势** | ①字节流量生态无可比拟②用户规模最大③产品迭代速度极快 |
| **劣势** | ①技术差异化不足②用户粘性依赖生态而非产品本身③模型能力非业界最强 |
| **机会** | ①企业级市场尚未深耕②海外市场潜力大（TikTok优势） |
| **威胁** | ①监管风险②DeepSeek低价冲击API业务③用户隐私争议 |

### Kimi SWOT

| | 内容 |
|--|------|
| **优势** | ①长文本技术壁垒②付费用户质量高③品牌口碑好 |
| **劣势** | ①用户规模相对小②缺乏流量生态③多模态能力较弱 |
| **机会** | ①深度研究场景扩张②企业知识库市场③海外扩张 |
| **威胁** | ①大厂复制长上下文能力②OpenAI直接竞争③融资压力 |

### DeepSeek SWOT

| | 内容 |
|--|------|
| **优势** | ①开源生态护城河②极致成本优势③全球开发者社区 |
| **劣势** | ①C端产品体验打磨不足②商业变现路径单一③品牌认知集中在技术圈 |
| **机会** | ①全球企业API市场②国产替代需求旺盛③开源生态持续扩大 |
| **威胁** | ①地缘政治风险（海外限制）②大厂开源跟进③算力资源限制 |

## 九、关键差异与壁垒（Moat Analysis）

三款产品的核心壁垒截然不同：

- **豆包的壁垒是流量分发**：字节系App日活超7亿，豆包天然获得用户触达优势，这是其他产品无法复制的
- **Kimi的壁垒是技术+用户质量**：长上下文技术积累形成认知护城河，高质量付费用户黏性强
- **DeepSeek的壁垒是开源生态**：一旦开发者基于其模型构建应用，迁移成本极高；开源模型的改进贡献也反哺自身

## 十、机会点与策略建议（Opportunities）

1. **垂直场景深耕**：三款产品均以通用能力为主，法律、医疗、教育、金融等垂直行业的专业化需求未被充分满足。建议新进入者选择1-2个高付费意愿的垂直场景切入，而非与巨头正面竞争。

2. **企业私有化部署**：数据安全顾虑使大量企业不愿使用公有云AI服务，私有化部署市场潜力巨大。DeepSeek开源模型已在这一方向获得先发优势，其他玩家可考虑混合云策略。

3. **AI Agent生态构建**：单次对话的竞争已趋于饱和，能够完成多步骤任务的Agent（如豆包的Coze）将成为下一轮竞争重点。先建立开发者生态的平台将获得显著先发优势。

## 十一、数据附录（Appendix）

**主要来源**
- 字节跳动2024Q4财报（豆包DAU数据）
- 月之暗面融资文件（Kimi付费用户数据）
- DeepSeek技术报告 arxiv.org/abs/2401.02954（训练成本数据）
- IDC《2024中国AI市场报告》（市场规模数据）
- QuestMobile《2025年Q1移动互联网报告》（行业MAU数据）
- 36kr《月之暗面2024年终盘点》
- 苹果App Store排行榜（DeepSeek下载数据）

**调研方法**：公开资料整理 + 官网信息核实 + 第三方数据交叉验证
"""

DEMO_REPORT_PM_TOOLS = """# 竞品调研报告：企业协作与项目管理SaaS

> **报告说明**：本报告聚焦中国企业协作市场，数据截至2025年Q1，来源包括官网、极客公园、ITValue、36kr等。

## 一、报告概述（Executive Summary）

企业协作与项目管理SaaS市场正经历「平台整合」与「AI化」双重变革。飞书、钉钉已从单一IM工具演进为「全家桶」平台，Jira则凭借专业项目管理能力在技术团队维持强势地位。AI能力正成为新的核心竞争维度，三款产品均在2024年大规模推进AI功能集成。

**核心结论：**
- 钉钉DAU超2亿，用户规模遥遥领先（阿里云2024年报）
- 飞书ARR超10亿元，是三者中商业化能力最强的（36kr，2024Q4）
- Jira在研发团队中渗透率达68%，专业壁垒最高（Atlassian年报）

## 二、市场与赛道分析（Market Context）

**市场规模与增速**

- 2024年中国企业协作SaaS市场规模约320亿元，同比增长28%（IDC）
- 预计2027年达650亿，CAGR约25%（艾瑞咨询）
- 中小企业付费意愿持续提升，ARPU同比增长15%

**竞争格局**

| 梯队 | 玩家 | 特点 |
|------|------|------|
| 综合平台 | 飞书、钉钉 | IM+文档+项目+OA全覆盖 |
| 专业工具 | Jira、Trello | 垂直深度，研发/PM首选 |
| 垂直新兴 | 线性、Notion | 小而美，设计/创意团队 |

**趋势判断**
- AI Copilot成为标配，会议纪要、文档摘要、任务自动拆解是高频场景
- 平台化竞争加剧，单点工具生存空间收窄
- 中小企业付费能力提升，但价格敏感度仍高，免费策略依然是拉新核心手段

## 三、竞品选择与分层（Competitive Landscape）

```
飞书   ──── 精品路线 ──── 高客单价，服务中大型企业
钉钉   ──── 规模路线 ──── 免费为主，覆盖全量中小企业
Jira   ──── 专业路线 ──── 研发团队刚需，技术壁垒最高
```

## 四、核心能力拆解（Product Capability Analysis）

### 飞书

- **产品定位**：面向中大型企业的一体化协作平台，强调"效率即文化"
- **核心功能**：飞书文档（多维表格）、飞书会议（AI纪要）、飞书项目、OKR管理、审批流、飞书智能伙伴
- **技术特点**：自研实时协作引擎，飞书文档性能业界领先；AI能力基于字节MegaScale
- **定价策略**：标准版30元/人/月；高级版60元/人/月；企业版定制（来源：[飞书官网定价](https://www.feishu.cn/price)）
- **用户规模**：企业客户数超50万，付费企业超10万（36kr，2024Q4）
- **更新节奏**：每月1-2次版本迭代，AI功能密集更新
- **分发渠道**：直销为主（销售团队），App/PC客户端，应用市场
- **商业模式**：订阅制SaaS；飞书应用市场生态分成；ISV合作
- **近期动态**：2024年发布「飞书智能伙伴」，深度集成AI；推出飞书People（HR模块）

### 钉钉

- **产品定位**：面向全量企业用户（重点是中小企业）的数字化工作平台
- **核心功能**：IM、视频会议、文档协作、钉钉低代码（宜搭）、钉钉AI助理、考勤打卡、审批
- **技术特点**：阿里云底层支撑，钉钉AI基于通义大模型；低代码能力是核心差异点
- **定价策略**：基础版永久免费；专业版9元/人/月；专属版30元/人/月（来源：[钉钉官网](https://www.dingtalk.com/pricing)）
- **用户规模**：DAU超2亿，企业组织数超2300万（阿里云2024年报）
- **更新节奏**：每季度大版本，月度小版本
- **分发渠道**：免费策略驱动下载，政务采购，教育行业合作
- **商业模式**：免费拉新→付费升级（功能限制）；阿里云绑定销售；行业解决方案
- **近期动态**：2024年推出「AI助理」全面覆盖工作场景；与阿里云深度绑定，主打AI+低代码

### Jira

- **产品定位**：全球最主流的敏捷项目管理工具，研发团队标配
- **核心功能**：Scrum/Kanban看板、Sprint规划、Bug追踪、发布管理、报表分析、Atlassian Intelligence（AI）
- **技术特点**：与Confluence、Bitbucket深度集成，形成研发全链路工具链；REST API生态成熟
- **定价策略**：Free（10人以内免费）；Standard $8.15/人/月；Premium $16/人/月（来源：[Jira官网](https://www.atlassian.com/software/jira/pricing)）
- **用户规模**：全球超30万企业客户，研发团队渗透率68%（Atlassian 2024年报）
- **更新节奏**：云版本持续交付，季度大版本
- **分发渠道**：自助注册，渠道合作伙伴（SI/ISV），Atlassian Marketplace
- **商业模式**：按用户订阅；Marketplace应用分成；企业级支持服务
- **近期动态**：2024年推出Jira AI，支持Epic/Story自动拆解、冲刺摘要；加速云化迁移

## 五、商业模式分析（Monetization）

| 维度 | 飞书 | 钉钉 | Jira |
|------|------|------|------|
| 核心模式 | 订阅制 | 免费+升级 | 订阅制 |
| 起步价 | 30元/人/月 | 免费 | $8.15/人/月 |
| 付费触发点 | 功能/存储限制 | 人数/功能限制 | 人数超10人 |
| ARPU | 高（>2000元/企业/年） | 低（<500元） | 中高（视规模） |
| 收入多样性 | 中（订阅为主） | 高（订阅+云+低代码） | 高（订阅+市场+服务） |

## 六、增长与分发策略（Growth Strategy）

**飞书**：高端品牌策略，主打"用飞书的公司都是好公司"，通过标杆客户（理想汽车、SHEIN等）形成背书效应；销售驱动，销售团队超3000人。

**钉钉**：免费策略+政务教育捆绑，在疫情期间"线上办公"热潮中完成用户基数积累；低代码生态吸引行业ISV共建解决方案。

**Jira**：Product-Led Growth (PLG)典范——个人/小团队免费使用，随组织规模扩大自然升级付费；Atlassian Marketplace超5000个应用形成强生态护城河。

## 七、用户与场景分析（User & Use Case）

| 用户群 | 飞书 | 钉钉 | Jira |
|--------|------|------|------|
| 核心用户 | 中大型企业、互联网公司 | 全量中小企业、传统企业 | 研发团队、技术公司 |
| 典型角色 | HR、PM、管理者 | 老板、HR、销售 | 工程师、Scrum Master |
| 核心场景 | 远程协作、OKR管理、文档共创 | 考勤管理、审批流、通知触达 | Sprint规划、Bug追踪、发布管理 |
| 决策者 | CTO/HR VP | 老板（中小企业主） | 研发总监/PM |

## 八、优劣势对比（SWOT / 对比矩阵）

### 飞书 SWOT

| | 内容 |
|--|------|
| **优势** | ①产品体验业界最佳②多维表格技术独特③字节文化背书吸引互联网企业 |
| **劣势** | ①价格偏高，中小企业转化难②国内市场份额不及钉钉③海外推广受限 |
| **机会** | ①AI Copilot深度整合②出海市场（Lark）③HR/财务模块扩展 |
| **威胁** | ①钉钉价格竞争②企业微信生态②Notion等小而美工具分流 |

### 钉钉 SWOT

| | 内容 |
|--|------|
| **优势** | ①用户规模压倒性领先②免费策略覆盖面广③阿里云绑定优势 |
| **劣势** | ①产品体验口碑不及飞书②ARPU极低③高端市场渗透难 |
| **机会** | ①低代码市场扩张②政务数字化持续投入③AI助理商业化 |
| **威胁** | ①飞书向下渗透②企业微信②政策风险 |

### Jira SWOT

| | 内容 |
|--|------|
| **优势** | ①研发团队中强认知壁垒②Atlassian生态协同③全球市场地位稳固 |
| **劣势** | ①界面复杂，学习曲线陡峭②价格对中小团队偏高③本土化服务弱 |
| **机会** | ①AI辅助项目管理②国内研发团队云化迁移③与Confluence深化整合 |
| **威胁** | ①Linear等新兴工具冲击②飞书项目功能加强③国产替代趋势 |

## 九、关键差异与壁垒（Moat Analysis）

- **飞书**：核心壁垒是产品体验 + 企业文化认同。飞书文档的多维表格已成为替代Excel的解决方案，形成数据迁移壁垒；字节的工程师文化吸引追求效率的互联网企业。

- **钉钉**：核心壁垒是用户规模 + 政务生态。超2300万企业组织的网络效应难以撼动；政务和教育市场的深度绑定形成强护城河。

- **Jira**：核心壁垒是工作流沉淀。企业在Jira中积累数年的Issue、Sprint、发布历史难以迁移；与CI/CD工具链的深度集成使替换成本极高。

## 十、机会点与策略建议（Opportunities）

1. **AI项目管理助手**：三款产品的AI功能仍以"锦上添花"为主（会议纪要、文档摘要），真正能自动拆解需求、分配任务、预测风险的AI项目管理能力尚属空白。先做到此能力的产品将获得显著优势。

2. **中小企业精品化**：钉钉占据中小企业规模，但产品体验粗糙；飞书产品好但价格高。"中小企业版飞书"——同等产品体验+更低价格——是一个明显的市场缺口。

3. **研发效能整合**：Jira + GitHub + Confluence的工具链割裂是研发团队的痛点，能真正打通"需求-开发-测试-发布"全链路的国产工具有机会借国产替代趋势获得突破。

## 十一、数据附录（Appendix）

**主要来源**
- 阿里云2024年报（钉钉DAU、企业数数据）
- 36kr《飞书2024年终盘点》（ARR数据）
- Atlassian FY2024 Annual Report（Jira全球客户数、渗透率）
- IDC《2024年中国企业协作市场报告》
- 飞书官网定价页 feishu.cn/price
- 钉钉官网定价页 dingtalk.com/pricing
- Jira官网定价页 atlassian.com/software/jira/pricing
"""

DEMO_REPORT_WRITING = """# 竞品调研报告：AI写作与知识管理工具

> **报告说明**：本报告聚焦AI写作与知识管理赛道，数据截至2025年Q1。来源包括官网、少数派、AppSo、Product Hunt等。

## 一、报告概述（Executive Summary）

AI写作工具正从"辅助工具"升级为"知识工作基础设施"。Notion AI凭借国际市场先发优势和高质量用户基础稳居领先；飞书AI依托企业协作场景形成场景壁垒；语雀则以阿里生态背书在国内知识管理赛道深耕。2025年，三款产品均面临如何将AI能力深度嵌入用户日常工作流而非停留在"功能展示"层面的核心挑战。

**核心结论：**
- Notion全球付费用户超400万，估值100亿美元（Notion融资公告，2023）
- 飞书文档DAU超1000万，AI功能月活渗透率35%（飞书2024年报）
- 语雀注册用户超500万，主要面向国内开发者和技术团队（语雀官网）

## 二、市场与赛道分析（Market Context）

**市场规模与增速**

- 2024年全球AI写作工具市场规模约60亿美元，同比增长180%（Grand View Research）
- 中国市场约15亿美元，增速超全球平均水平
- 知识管理SaaS与AI写作工具边界模糊，正加速融合

**竞争格局**

| 类型 | 代表产品 | 核心诉求 |
|------|---------|---------|
| 国际全能型 | Notion AI | 灵活度最高，个人+团队 |
| 企业协作型 | 飞书AI | 场景嵌入，大企业优先 |
| 国内垂直型 | 语雀、石墨 | 文档协作+知识沉淀 |

**趋势判断**
- 单纯"AI改写/润色"已成标配，差异化竞争转向"知识连接"（让AI理解你的所有文档）
- 团队知识库 + AI问答成为企业级核心需求
- 个人知识管理（PKM）与团队协作工具边界模糊化

## 三、竞品选择与分层（Competitive Landscape）

```
Notion AI  ── 国际标杆 ── 高灵活度，创意工作者+技术团队
飞书AI     ── 企业嵌入 ── 场景化，中大型企业协作
语雀       ── 国内专注 ── 文档+知识库，开发者/技术团队
```

## 四、核心能力拆解（Product Capability Analysis）

### Notion AI

- **产品定位**：全球最具影响力的个人&团队知识管理工具，AI是深度嵌入的核心能力
- **核心功能**：块编辑器、数据库、AI写作助手、AI问答（Ask AI）、模板市场、协作评论
- **技术特点**：基于GPT-4o，AI能理解整个Workspace的上下文；块状编辑体验业界独特
- **定价策略**：Free版有限制；Plus $12/月；Business $18/人/月；AI功能额外+$10/月（来源：[Notion官网定价](https://www.notion.so/pricing)）
- **用户规模**：全球注册用户3000万+，付费用户400万+（Notion 2023融资披露）
- **更新节奏**：每月1-2次功能更新，AI能力持续迭代
- **分发渠道**：自助注册（PLG），教育免费计划，企业销售
- **商业模式**：个人订阅 + 团队订阅 + 企业版 + AI附加包
- **近期动态**：2024年推出Notion AI Q&A（全库问答），支持跨页面语义搜索；推出Notion Sites

### 飞书AI

- **产品定位**：嵌入飞书生态的AI写作与知识管理能力，强调"会议-文档-任务"一体化
- **核心功能**：飞书文档（实时协作）、多维表格、飞书知识库、AI起草/润色/摘要、会议AI纪要、飞书智能伙伴
- **技术特点**：字节自研模型+接入主流大模型；多维表格的实时协作性能行业领先
- **定价策略**：飞书标准版含基础AI能力；AI高级功能在飞书智能伙伴中，30元+/人/月（来源：[飞书官网](https://www.feishu.cn/price)）
- **用户规模**：飞书文档DAU超1000万，AI月活渗透35%（飞书2024年报）
- **更新节奏**：AI功能每月迭代，文档编辑器每季度大更新
- **分发渠道**：飞书App内原生，无需额外下载；企业销售驱动
- **商业模式**：飞书整体订阅包含（无需单独购买AI）；高级AI能力按模块收费
- **近期动态**：2024年发布"飞书智能伙伴"，支持自定义AI Bot；知识库AI问答上线

### 语雀

- **产品定位**：面向国内开发者和技术团队的专业知识管理工具
- **核心功能**：文档编辑（支持Markdown）、知识库、团队空间、语雀画板、AI辅助写作
- **技术特点**：基于阿里云基础设施，接入通义千问；Markdown渲染质量高，代码块友好
- **定价策略**：个人版免费；团队版89元/人/年；企业版定制（来源：[语雀官网定价](https://www.yuque.com/settings/billing)）
- **用户规模**：注册用户500万+，主要集中在技术团队（语雀官网披露）
- **更新节奏**：每季度大版本，AI能力相对保守
- **分发渠道**：阿里系生态（钉钉内置）、独立App、开发者社区口碑
- **商业模式**：订阅制；与阿里云/钉钉捆绑销售
- **近期动态**：2024年接入通义AI，推出AI辅助写作；小程序版本上线

## 五、商业模式分析（Monetization）

| 维度 | Notion AI | 飞书AI | 语雀 |
|------|-----------|--------|------|
| 核心模式 | 订阅（个人+团队）+AI附加 | 整体订阅包含 | 订阅（低价） |
| 个人价格 | $10-22/月 | 30元/月起 | 89元/年 |
| AI能力定价 | 单独+$10/月 | 含在套餐内 | 含在套餐内 |
| 主要收入 | 付费订阅 | 企业合同 | 企业合同 |
| 变现成熟度 | 高 | 中 | 低 |

## 六、增长与分发策略（Growth Strategy）

**Notion AI**：PLG经典案例——个人免费使用，团队协作触发付费升级；模板市场形成创作者生态；Reddit/Twitter口碑传播是重要获客渠道；中国市场因访问限制增长受阻。

**飞书AI**：依托飞书企业销售体系推送，AI能力作为飞书产品的增值点而非独立产品售卖；飞书开放平台吸引ISV构建AI应用。

**语雀**：阿里/钉钉生态自然导流；开发者社区口碑（程序员圈子"人手一个语雀"）；少数派等媒体报道贡献质量流量。

## 七、用户与场景分析（User & Use Case）

| 维度 | Notion AI | 飞书AI | 语雀 |
|------|-----------|--------|------|
| 核心用户 | 创意工作者、PM、设计师、学生 | 企业员工、管理者 | 技术团队、开发者 |
| 核心场景 | 个人知识管理、项目规划、内容创作 | 会议纪要、文档协作、知识库 | 技术文档、API文档、团队Wiki |
| 使用驱动 | 自驱（个人习惯）| 被动（企业要求） | 自驱（技术需求） |
| 付费决策者 | 个人 | 公司 | 团队/个人 |

## 八、优劣势对比（SWOT / 对比矩阵）

### Notion AI SWOT

| | 内容 |
|--|------|
| **优势** | ①产品灵活度最高②全球用户基础③AI与内容深度融合 |
| **劣势** | ①国内访问受限②学习成本高③协作性能不及飞书 |
| **机会** | ①企业级市场深耕②更多语言本土化③API生态扩展 |
| **威胁** | ①飞书国内竞争②Linear等专业工具分流③AI写作工具同质化 |

### 飞书AI SWOT

| | 内容 |
|--|------|
| **优势** | ①场景嵌入深②企业客户黏性高③字节技术实力背书 |
| **劣势** | ①个人用户体验弱②AI能力非核心卖点③灵活度不及Notion |
| **机会** | ①智能伙伴商业化②出海市场（Lark）③知识库AI问答深化 |
| **威胁** | ①钉钉低价竞争②Notion国内回流③企业微信 |

### 语雀 SWOT

| | 内容 |
|--|------|
| **优势** | ①技术团队中口碑极好②Markdown体验最佳③价格亲民 |
| **劣势** | ①用户规模小④AI能力迭代慢③非阿里员工感知度低 |
| **机会** | ①开发者生态扩展②AI知识库问答③钉钉深度整合 |
| **威胁** | ①飞书文档分流②Notion中国用户回流③Obsidian等本地工具 |

## 九、关键差异与壁垒（Moat Analysis）

- **Notion**：壁垒是灵活性 + 用户习惯。用户在Notion中搭建的复杂数据库、工作流迁移成本极高；全球社区生态（模板、插件）持续强化网络效应。

- **飞书AI**：壁垒是场景嵌入。AI能力在会议→纪要→任务→文档的完整链路中无缝流转，单独工具无法复制这种体验；企业合同锁定带来的切换成本高。

- **语雀**：壁垒是技术社区认同。在程序员圈子中形成"专业文档用语雀"的心智，技术文档的SEO积累也形成一定护城河。

## 十、机会点与策略建议（Opportunities）

1. **AI知识连接**：三款产品的AI仍停留在"单文档处理"，真正能跨文档理解、连接知识图谱的能力尚不成熟。"让AI真正读懂你所有的笔记和文档"是下一个杀手级功能。

2. **垂直行业模板生态**：Notion的模板市场已证明创作者经济模式有效，国内市场尚无类似平台。飞书或语雀若开放模板市场并与创作者分成，有望形成强生态壁垒。

3. **个人知识管理 + 团队协作打通**：目前个人笔记和团队文档是割裂的（个人用Notion/Obsidian，团队用飞书），能同时满足"个人深度思考"和"团队高效协作"的产品是空白。

## 十一、数据附录（Appendix）

**主要来源**
- Notion 2023年融资公告（估值、用户数数据）
- 飞书2024年报（文档DAU、AI渗透率）
- 语雀官网披露信息（用户规模）
- Grand View Research《AI写作工具市场报告2024》
- Notion官网定价页 notion.so/pricing
- 飞书官网定价页 feishu.cn/price
- 语雀官网定价页 yuque.com/settings/billing
"""

# ── 静态报告映射 ──────────────────────────────────────────────────────────────
STATIC_REPORTS = {
    "ai_chat": DEMO_REPORT_AI_CHAT,
    "pm_tools": DEMO_REPORT_PM_TOOLS,
    "writing": DEMO_REPORT_WRITING,
}

# ── Demo 报告配置 ─────────────────────────────────────────────────────────────
DEMO_REPORTS = {
    "ai_chat": {
        "track": "AI 对话产品",
        "title": "豆包 vs Kimi vs DeepSeek",
        "companies": "豆包,Kimi,DeepSeek",
        "market": "国内AI对话助手",
    },
    "writing": {
        "track": "AI 写作工具",
        "title": "Notion AI vs 飞书AI vs 语雀",
        "companies": "Notion AI,飞书AI,语雀",
        "market": "AI写作与知识管理工具",
    },
    "pm_tools": {
        "track": "项目管理 SaaS",
        "title": "飞书 vs 钉钉 vs Jira",
        "companies": "飞书,钉钉,Jira",
        "market": "企业协作与项目管理SaaS",
    },
}

# ── Word 下载功能 ─────────────────────────────────────────────────────────────
def markdown_to_docx_bytes(markdown_text: str, title: str) -> bytes:
    """把Markdown转成Word文件bytes，无需pandoc"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()

    # 设置页边距
    for section in doc.sections:
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    lines = markdown_text.split("\n")

    for line in lines:
        line = line.rstrip()

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x0f, 0x34, 0x60)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.runs[0].font.size = Pt(10.5)
        elif line.startswith("|"):
            # 简单处理表格行（跳过分隔行）
            if re.match(r"^\|[\s\-\|]+\|$", line):
                continue
            cells = [c.strip() for c in line.strip("|").split("|")]
            if not hasattr(doc, "_current_table"):
                doc._current_table = None
            if doc._current_table is None:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                doc._current_table = table
                row = table.rows[0]
                for i, cell in enumerate(cells):
                    row.cells[i].text = cell
                    row.cells[i].paragraphs[0].runs[0].bold = True
            else:
                row = doc._current_table.add_row()
                for i, cell in enumerate(cells[:len(row.cells)]):
                    row.cells[i].text = cell
        elif line.startswith(">"):
            p = doc.add_paragraph(line[1:].strip())
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x92, 0xa4)
        elif line.strip() == "" or line.strip() == "---":
            if hasattr(doc, "_current_table"):
                doc._current_table = None
            if line.strip() == "":
                pass
        elif line.startswith("```"):
            pass
        else:
            if hasattr(doc, "_current_table") and doc._current_table:
                doc._current_table = None
            # 处理行内粗体 **text**
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", clean)
            p = doc.add_paragraph(clean)
            p.runs[0].font.size = Pt(10.5) if p.runs else None

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Agent 工具 ────────────────────────────────────────────────────────────────
@tool
def search_web(query: str) -> str:
    """搜索网络，获取竞争对手的公开信息。输入搜索词，返回多条搜索结果摘要。"""
    try:
        try:
            from ddgs import DDGS
        except ImportError:
            from duckduckgo_search import DDGS
        results = list(DDGS().text(query, max_results=5))
        if not results:
            return "未找到搜索结果，请换一个更具体的搜索词。"
        lines = []
        for r in results:
            title = r.get("title", "").strip()
            body = r.get("body", "").strip()
            href = r.get("href", "").strip()
            lines.append(f"### {title}\n{body}\n来源: {href}")
        return "\n\n---\n\n".join(lines)
    except Exception as e:
        return f"搜索失败: {str(e)}"


@tool
def fetch_webpage(url: str) -> str:
    """抓取指定网页并提取纯文本内容（最多返回3000字符）。"""
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"
        ),
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        html = resp.text
        html = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<style[^>]*>.*?</style>", " ", html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r"<[^>]+>", " ", html)
        html = re.sub(r"&[a-z]+;", " ", html)
        html = re.sub(r"\s{2,}", "\n", html).strip()
        return textwrap.shorten(html, width=3000, placeholder="…（内容已截断）")
    except requests.exceptions.Timeout:
        return f"请求超时: {url}"
    except Exception as e:
        return f"无法获取页面: {str(e)}"


# ── 加载 Skills ──────────────────────────────────────────────────────────────
def load_skills():
    """加载所有 PM 技能库"""
    import os
    skills_content = []
    skills_dir = "skills"

    # 加载 claude-skills 的竞品分析
    claude_skill_path = os.path.join(skills_dir, "claude-skills/competitive-analysis/SKILL.md")
    if os.path.exists(claude_skill_path):
        with open(claude_skill_path, encoding="utf-8") as f:
            skills_content.append(f.read())

    # 加载 lenny-pm-skills 中与竞品分析相关的核心技能
    lenny_skills_dir = os.path.join(skills_dir, "lenny-pm-skills")
    if os.path.exists(lenny_skills_dir):
        # 只加载与竞品分析最相关的技能，避免 prompt 过长
        relevant_skills = [
            "09-竞品分析-competitive-analysis",
            "60-定位传播-positioning",
            "04-产品愿景-product-vision",
            "05-北极星指标-north-star-metrics",
            "07-路线图优先级-roadmap-prioritization",
        ]
        for skill_name in relevant_skills:
            skill_path = os.path.join(lenny_skills_dir, skill_name, "SKILL.md")
            if os.path.exists(skill_path):
                with open(skill_path, encoding="utf-8") as f:
                    skills_content.append(f.read())

    return "\n\n---\n\n".join(skills_content) if skills_content else ""


# ── Agent Prompt ──────────────────────────────────────────────────────────────
SKILLS_LIBRARY = load_skills()

SYSTEM_PROMPT = f"""你是一名专业的竞争情报（Competitive Intelligence）分析师。

## 方法论库

你可以参考以下产品经理专业方法论来指导你的分析：

{SKILLS_LIBRARY}

---

## 工作流程

阶段一：搜集竞品信息（每个竞品搜索2次 + 抓取1次官网）
阶段二：直接输出完整报告，不要输出过渡句

来源标注规则（严格执行）：
- 所有来源链接必须如实标注来源类型，例如：来源：搜狐科技报道、来源：36kr评测、来源：官网定价页
- 禁止将第三方媒体文章（搜狐、知乎、36kr、微信公众号等）标注为"官网"
- 官网链接仅限于该产品的官方域名
- 每条关键结论后附来源，格式：（来源：[标题](URL)）

报告格式要求：
- 以 "# 竞品调研报告：{{市场}}" 开头
- 包含以下11个章节，每章节内容必须实质性、有数据支撑：

  ## 一、报告概述（Executive Summary）
  ## 二、市场与赛道分析（Market Context）
  （至少包含：市场规模、增速、主要玩家、竞争格局、趋势判断，≥3条要点）
  ## 三、竞品选择与分层（Competitive Landscape）
  ## 四、核心能力拆解（Product Capability Analysis）
  （每个竞品用 ### 单独列小节，必须包含以下9个字段：
    产品定位、核心功能、技术特点、定价策略、用户规模、更新节奏、分发渠道、商业模式、近期动态）
  ## 五、商业模式分析（Monetization）
  （至少包含：收费方式、客单价区间、付费转化路径、主要收入来源，≥3条要点）
  ## 六、增长与分发策略（Growth Strategy）
  ## 七、用户与场景分析（User & Use Case）
  ## 八、优劣势对比（SWOT / 对比矩阵）
  （每个竞品各自列出：优势≥3条、劣势≥3条、机会≥2条、威胁≥2条；附整体对比矩阵表格）
  ## 九、关键差异与壁垒（Moat Analysis）
  ## 十、机会点与策略建议（Opportunities）
  （≥3条具体可执行的建议，每条说明依据）
  ## 十一、数据附录（Appendix）
  （列出本报告引用的所有来源URL，按章节归类）

- 使用表格对比关键数据
- 总字数不少于4000字
- 报告须有深度，避免泛泛而谈，每个判断都要有数据或事实支撑"""


# ── 构建 Agent ────────────────────────────────────────────────────────────────
class AgentState(TypedDict):
    messages: Annotated[list[BaseMessage], add_messages]


def build_agent(api_key: str, base_url: str, model: str):
    tools = [search_web, fetch_webpage]
    llm = ChatOpenAI(
        model=model,
        api_key=api_key,
        base_url=base_url,
        temperature=0.3,
        max_tokens=4096,
    ).bind_tools(tools)
    tool_node = ToolNode(tools)
    MAX_TOOL_CALLS = 50

    def agent_node(state: AgentState) -> dict:
        messages = list(state["messages"])
        if not any(isinstance(m, SystemMessage) for m in messages):
            messages = [SystemMessage(content=SYSTEM_PROMPT)] + messages
        tool_call_count = sum(
            1 for m in messages
            if isinstance(m, AIMessage) and hasattr(m, "tool_calls") and m.tool_calls
        )
        if tool_call_count >= MAX_TOOL_CALLS:
            messages.append(SystemMessage(
                content="信息已足够，请立即输出完整竞争情报报告，不要再调用任何工具。"
            ))
        response = llm.invoke(messages)
        return {"messages": [response]}

    def route(state: AgentState) -> Literal["tools", "agent", "__end__"]:
        last: AIMessage = state["messages"][-1]
        if hasattr(last, "tool_calls") and last.tool_calls:
            tool_call_count = sum(
                1 for m in state["messages"]
                if isinstance(m, AIMessage) and hasattr(m, "tool_calls") and m.tool_calls
            )
            if tool_call_count >= MAX_TOOL_CALLS:
                return END
            return "tools"
        content = getattr(last, "content", "") or ""
        if "# 竞品调研报告" in content or "# 竞争情报报告" in content:
            return END
        return "agent"

    graph = StateGraph(AgentState)
    graph.add_node("agent", agent_node)
    graph.add_node("tools", tool_node)
    graph.set_entry_point("agent")
    graph.add_conditional_edges(
        "agent", route,
        {"tools": "tools", "agent": "agent", "__end__": END}
    )
    graph.add_edge("tools", "agent")
    return graph.compile()


# ── 运行 Agent 并实时展示 ─────────────────────────────────────────────────────
def run_agent_with_ui(competitors: str, market: str, api_key: str, base_url: str, model: str):
    compiled = build_agent(api_key, base_url, model)
    user_task = (
        f"请对以下竞争对手进行全面研究，生成竞争情报报告：\n\n"
        f"**竞争对手：** {competitors}\n"
        f"**市场/产品类别：** {market}\n"
        f"**地理范围：** 中国市场为主\n\n"
        f"请先用工具搜集每个竞品的信息，搜集完毕后输出完整报告。"
    )
    log_placeholder = st.empty()
    report_placeholder = st.empty()
    log_lines = [
        "▶ Agent 启动",
        f"◆ 竞品：{competitors}",
        f"◆ 赛道：{market}",
        "─" * 38,
    ]
    tool_count = 0
    final_report = ""

    def update_log():
        display = "\n".join(log_lines[-18:])
        log_placeholder.markdown(
            f'<div class="log-box">{display}</div>',
            unsafe_allow_html=True
        )

    update_log()
    try:
        for chunk in compiled.stream(
            {"messages": [("user", user_task)]},
            stream_mode="updates",
        ):
            for node_name, node_output in chunk.items():
                msgs: list[BaseMessage] = node_output.get("messages", [])
                if node_name == "agent":
                    for msg in msgs:
                        if not isinstance(msg, AIMessage):
                            continue
                        if hasattr(msg, "tool_calls") and msg.tool_calls:
                            for tc in msg.tool_calls:
                                tool_count += 1
                                args = tc.get("args", {})
                                arg_val = next(iter(args.values()), "") if args else ""
                                preview = str(arg_val)[:55]
                                icon = "🔍" if tc["name"] == "search_web" else "🌐"
                                log_lines.append(f"[{tool_count:02d}] {icon} {tc['name']}: {preview}")
                                update_log()
                        elif msg.content and (
                            "# 竞品调研报告" in msg.content or "# 竞争情报报告" in msg.content
                        ):
                            final_report = msg.content
                            log_lines.append("─" * 38)
                            log_lines.append("✅ 报告生成完成！")
                            update_log()
                elif node_name == "tools":
                    for msg in msgs:
                        content = getattr(msg, "content", "") or ""
                        preview = content[:45].replace("\n", " ")
                        log_lines.append(f"   ↳ {preview}...")
                    update_log()
    except Exception as e:
        st.error(f"Agent运行出错：{str(e)}")
        return ""

    if final_report:
        log_placeholder.empty()
        st.success(f"✅ 完成！Agent共调用工具 {tool_count} 次")
        return final_report
    else:
        st.warning("⚠️ 未能生成完整报告，请重试")
        return ""


# ── 显示报告（含下载按钮）────────────────────────────────────────────────────
def show_report(content: str, title: str):
    """展示报告 + Word下载按钮"""
    # 修复标题中的链接（Streamlit不渲染标题里的链接，改为正文）
    fixed = re.sub(
        r"^(#{1,4} .+)\[([^\]]+)\]\(([^\)]+)\)",
        r"\1\2（\3）",
        content,
        flags=re.MULTILINE
    )

    col_report, col_dl = st.columns([5, 1])
    with col_dl:
        try:
            import docx  # noqa
            docx_bytes = markdown_to_docx_bytes(content, title)
            if docx_bytes:
                st.download_button(
                    label="⬇️ 下载 Word",
                    data=docx_bytes,
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        except ImportError:
            st.caption("安装 python-docx 可下载Word版本")

    st.markdown(f'<div class="report-body">', unsafe_allow_html=True)
    st.markdown(fixed)
    st.markdown("</div>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# 页面布局
# ══════════════════════════════════════════════════════════════════════════════

# ── 侧边栏：仅放 API 配置 ────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ API 配置")
    st.caption("自定义分析需要填写，Demo报告无需填写")
    st.markdown("---")

    default_key = ""
    default_url = ""
    default_model = ""
    try:
        default_key = st.secrets.get("API_KEY", "")
        default_url = st.secrets.get("API_BASE_URL", "")
        default_model = st.secrets.get("MODEL", "")
    except Exception:
        pass

    api_key_input = st.text_input("API Key", type="password", value=default_key, placeholder="sk-...")
    base_url_input = st.text_input("Base URL", value=default_url, placeholder="https://api.openai.com/v1")
    model_input = st.text_input("模型名称", value=default_model, placeholder="例：gpt-4o、deepseek-chat")

    st.markdown("---")
    st.markdown("""
    <div style='font-size:0.72rem; color:#8892a4; line-height:1.9'>
    <b style='color:#e94560'>技术栈</b><br>
    LangGraph ReAct Agent<br>
    DuckDuckGo 实时搜索<br>
    网页内容抓取<br>
    OpenAI 兼容接口<br><br>
    <b style='color:#e94560'>报告结构</b><br>
    11章 · 4000字+ · 附来源<br><br>
    <b style='color:#e94560'>Demo报告</b><br>
    预置3份，无需API Key<br>
    直接点击查看
    </div>
    """, unsafe_allow_html=True)


# ── Hero ──────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="hero">
    <h1>竞品分析 Agent <span class="badge">LIVE AGENT</span></h1>
    <div class="subtitle">真实联网搜索 · 网页抓取 · LangGraph ReAct · 11章专业报告</div>
</div>
""", unsafe_allow_html=True)

# 强制滚到顶部
import streamlit.components.v1 as components
components.html("<script>window.parent.document.querySelector('.main').scrollTo(0, 0);</script>", height=0)

# ── 主区域：标签页布局 ────────────────────────────────────────────────────────
tab_demo, tab_custom = st.tabs(["📄 Demo 报告（即点即看）", "🚀 自定义分析（输入你的竞品）"])

# ══ Tab 1：Demo 报告 ══════════════════════════════════════════════════════════
with tab_demo:
    st.markdown("#### 预置竞品分析报告")
    st.caption("以下报告已预先生成，无需 API Key，点击卡片直接查看完整内容 👇")

    # 三个卡片按钮
    col1, col2, col3 = st.columns(3)
    selected = st.session_state.get("demo_selected", "ai_chat")

    for col, (key, report) in zip([col1, col2, col3], DEMO_REPORTS.items()):
        with col:
            is_active = selected == key
            status = "✓ 当前查看" if is_active else ""
            active_style = "background:#fffbf0 !important; border-left-color:#c9a84c !important;" if is_active else ""
            st.markdown(f"""
            <style>
            div[data-testid="stButton"] button[kind="secondary"]#btn_{key} {{
                {active_style}
            }}
            </style>
            <div class="demo-card {'active' if is_active else ''}">
                <div class="track">{report['track']}</div>
                <div class="title">{report['title']}</div>
                <div class="status">{status}</div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("点击查看" if not is_active else "当前查看", key=f"demo_btn_{key}", use_container_width=True, type="secondary"):
                st.session_state["demo_selected"] = key
                st.rerun()

    st.markdown("---")

    # 展示当前选中的报告
    current = st.session_state.get("demo_selected", "ai_chat")
    report_info = DEMO_REPORTS[current]
    report_content = STATIC_REPORTS[current]

    st.markdown(f"### 📊 {report_info['title']}")
    st.caption(f"赛道：{report_info['track']} · 竞品：{report_info['companies'].replace(',', ' · ')}")

    show_report(report_content, report_info['title'])


# ══ Tab 2：自定义分析 ══════════════════════════════════════════════════════════
with tab_custom:
    st.markdown("#### 分析你指定的竞品")
    st.caption("填入竞品和赛道，Agent 将实时联网搜索并生成专属报告（约需 2-3 分钟）")

    st.markdown("**输入竞品信息**")

    custom_competitors = st.text_input(
        "竞品名称（英文逗号分隔）",
        placeholder="例：微信，钉钉，飞书",
        key="custom_comp_input"
    )
    custom_market = st.text_input(
        "所在赛道",
        placeholder="例：企业即时通讯",
        key="custom_market_input"
    )

    api_ok = bool(api_key_input)
    if not api_ok:
        st.warning("⚠️ 还没有配置 API Key — 点击页面左上角 > 展开侧边栏填写后即可开始分析")

    if st.button("🚀 开始分析", disabled=not api_ok, use_container_width=True):
        if not custom_competitors or not custom_market:
            st.error("请填写竞品名称和赛道")
        else:
            st.session_state["custom_run"] = {
                "competitors": custom_competitors,
                "market": custom_market,
                "api_key": api_key_input,
                "base_url": base_url_input,
                "model": model_input,
            }
            cache_key = f"custom_{custom_competitors}_{custom_market}"
            if cache_key in st.session_state:
                del st.session_state[cache_key]
            st.rerun()

    # 运行结果
    run_cfg = st.session_state.get("custom_run")
    if run_cfg:
        cache_key = f"custom_{run_cfg['competitors']}_{run_cfg['market']}"
        st.markdown("---")
        st.markdown(f"### 📊 自定义分析：{run_cfg['competitors']}")
        st.caption(f"赛道：{run_cfg['market']}")

        if st.session_state.get(cache_key):
            st.success("✅ 已缓存（本次会话有效）")
            show_report(st.session_state[cache_key], run_cfg['competitors'])
        else:
            result = run_agent_with_ui(
                competitors=run_cfg["competitors"],
                market=run_cfg["market"],
                api_key=run_cfg["api_key"],
                base_url=run_cfg["base_url"],
                model=run_cfg["model"],
            )
            if result:
                st.session_state[cache_key] = result
                show_report(result, run_cfg['competitors'])
